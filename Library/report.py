import base64
import time
import os
from bs4 import BeautifulSoup
import getpass
import xml.etree.ElementTree as ET
from docx import Document
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email.mime.text import MIMEText
from email import encoders
import smtplib
import matplotlib.pyplot as plt
from pdf2docx import Converter
from datetime import datetime, timedelta

class report(object):

    def read_word_document_content(self, file_path):
        doc = Document(file_path)
        content = ""
        for paragraph in doc.paragraphs:
            content += paragraph.text + "\n"
        return content
    
    def send_mail(self, sender, to, subject, body, attachments, host, email_username, email_password):
        msg = MIMEMultipart()
        msg['Subject'] = subject
        # msg['Body'] = body
        msg.attach(MIMEText(body, 'plain'))
        msg['From'] = sender
        # msg['To'] = to
        msg['To'] = ', '.join(to)  # Join multiple recipients with a comma

        for attachment in attachments:
            part = MIMEBase('application', "octet-stream")
            part.set_payload(open(attachment, "rb").read())
            encoders.encode_base64(part)
            part.add_header('Content-Disposition', 'attachment', filename=attachment)
            msg.attach(part)

        with smtplib.SMTP(host,587) as server:
            server.ehlo()
            server.starttls()
            server.ehlo()

            server.login(email_username,email_password)
            server.send_message(msg)

    def read_specific_data_from_html_and_write_to_file(self, url, xml_path, html_file_path, output_file_path, execution_mode='Remote'):
            current_username = getpass.getuser()
            if execution_mode=='Remote':
                start_time, end_time, execution_time, passed, failed, skipped, build_no, total_tc, suite_statistic = self.read_xml(xml_path)
            else:
                start_time, end_time, execution_time, passed, failed, skipped, build_no, total_tc, suite_statistic = self.read_local_xmlfile(xml_path)
            total_tc = int(passed)+int(failed)
            percentage = (int(passed)/int(total_tc))*100
            rounded_percentage = round(percentage, 2)
            # Read the HTML file
            with open(html_file_path, 'r', encoding='utf-8') as file:
                html_content = file.read()
            # Parse the HTML content
            soup = BeautifulSoup(html_content, 'html.parser')

            # Example: Extract data using Css
            time = 'a:-soup-contains("Execution Time(m):")'
            execution_time = soup.select_one(time)
            result = 'span:-soup-contains("Suite Statistics:") ~ div.col-md-12 table td.fail'
            execution_result = soup.select_one(result)

            if execution_time:
                specific_data = execution_time.get_text(strip=True)
                specific_data2 = execution_result.get_text(strip=True)
                result_text = "Result:  FAILURE" if int(specific_data2) > 0 else "Result:  SUCCESS"
                # Write to a text file
                with open(output_file_path + '.txt', 'w', encoding='utf-8') as txt_file:
                    txt_file.write(f"OVERALL SUMMARY\n\n")
                    txt_file.write(f"Triggered By:  {current_username}\n")
                    # txt_file.write(f"Environment Url:  {url}\n")
                    txt_file.write(f"Start Time:  {start_time}\n")
                    txt_file.write(f"End Time:  {end_time}\n")
                    txt_file.write(f"{specific_data}\n")
                    txt_file.write(f"{result_text}\n")
                    txt_file.write(f"Build Number:  {build_no}\n\n\n")
                    txt_file.write(f"TEST RESULTS\n\n")
                    txt_file.write(f"Passed: {passed}\n")
                    txt_file.write(f"Failed: {failed}\n")
                    txt_file.write(f"Total: {total_tc}\n")
                    txt_file.write(f"Pass Percentage: {rounded_percentage}%\n")

                # Write to a Word document (docx)
                # doc = Document()
                # doc.add_paragraph(f"OVERALL SUMMARY\n")
                # doc.add_paragraph(f"Triggered By:  {current_username}")
                # doc.add_paragraph(f"Environment Url:  {url}")
                # doc.add_paragraph(f"Start Time:  {start_time}")
                # doc.add_paragraph(f"End Time:  {end_time}")
                # doc.add_paragraph(f"{specific_data}")
                # doc.add_paragraph(f"{result_text}\n\n")
                # doc.add_paragraph(f"TEST RESULTS\n")
                # doc.add_paragraph(f"Total: {total_tc}")
                # doc.add_paragraph(f"Passed: {passed}")
                # doc.add_paragraph(f"Failed: {failed}")
                # doc.add_paragraph(f"Pass Percentage: {rounded_percentage}%")
                # doc.save(output_file_path + '.docx')

            else:
                print("Specific element not found.")

    def pdf_to_word(self, input_pdf, output_word):
        cv = Converter(input_pdf)
        cv.convert(output_word, start=0, end=None)
        cv.close()
        
    def read_xml(self, input_xml_path):
            # Parse the XML data from the file
            tree = ET.parse(input_xml_path)
            root = tree.getroot()
            status_elements = root.findall(".//test/status")
            # Initialize a counter for 'SKIP'
            skip_count = 0
            # Iterate through the status elements
            for status_element in status_elements:
                status_value = status_element.get('status')
                if status_value == 'SKIP':
                    skip_count += 1
            # Print the count of 'SKIP'
            print("Number of 'SKIP' status:", skip_count)
            statistics = root.find(".//statistics/total/stat")
            passed = statistics.get('pass')
            failed = statistics.get('fail')
            total_tc = int(passed)+int(failed)+int(skip_count)  
            print(passed, failed) 
            suite_statistics = root.findall(".//statistics/suite/stat")
            suite_statistic = []
            for index, suite in enumerate(suite_statistics):
                if index == 0:
                    # Skip the first iteration
                    continue
                suite_name = suite.get('name')
                suite_pass = suite.get('pass')
                suite_fail = suite.get('fail')
                suite_total = int(suite_pass)+int(suite_fail)
                suite_status = 'FAIL' if int(suite_fail) > 0 else 'PASS'
                suite_statistic.append([index, suite_name, suite_status, suite_total, suite_pass, suite_fail])
            print(suite_statistic)
            suite_with_stat = root.find(".//suite/status")
            if suite_with_stat is not None:
                starttime_value = suite_with_stat.get('start')
                elapsed_value = suite_with_stat.get('elapsed')
                start_datetime = datetime.fromisoformat(starttime_value)
                print(start_datetime)
                # Convert the elapsed time to a float
                elapsed_seconds = float(elapsed_value)
                execution_time = self.seconds_to_minutes_and_seconds(elapsed_seconds)
                print(elapsed_seconds)
                # Create a timedelta object with the elapsed time
                elapsed_timedelta = timedelta(seconds=elapsed_seconds)
                print(elapsed_timedelta)
                # Add the elapsed time to the start datetime
                end_datetime = start_datetime + elapsed_timedelta
                # Convert the result to a string
                endtime_value = end_datetime.isoformat()
                end_datetime = datetime.fromisoformat(endtime_value)
                start_time = start_datetime.replace(microsecond=0)
                start_time = str(start_time).replace("-", "")
                end_time = end_datetime.replace(microsecond=0)
                end_time = str(end_time).replace("-", "")
                print("Start Time:", start_time)
                print("Elapsed Time:", execution_time)
                print("End Time:", end_time)
            splitted_starttime = str(start_time).split(" ")
            date = splitted_starttime[0].replace("-", "")
            log_hr = splitted_starttime[1].split(":")[0]
            build_no = date+"."+log_hr
            print(build_no)
            return start_time, end_time, execution_time, passed, failed, skip_count, build_no, total_tc, suite_statistic
    
    def read_log_from_xml(self, input_xml_path):
        tree = ET.parse(input_xml_path)
        root = tree.getroot()
        suite_datas = root.findall(".//suite[@id='s1']")
        if suite_datas:
            first_suite = suite_datas[0]
            execution_name = first_suite.get("name")
            print("Execution Name:", execution_name)
            if execution_name == 'VLMS4.2 Regression':
                suite_datas = first_suite.findall(".//suite")
        data_rows = []
        suite_times = []
        for suite_data in suite_datas:
            suite_name = suite_data.get("name")
            suite_status = suite_data.find("./status")
            suite_time_in_seconds = suite_status.get("elapsed")
            suite_time = self.seconds_to_minutes_and_seconds(suite_time_in_seconds)
            suite_times.append(suite_time)
            test_datas= suite_data.findall("test")
            for test_data in test_datas:
                testcase_name = test_data.get("name")
                status_element = test_data.find("status")
                status = status_element.get("status")
                starttime_value = status_element.get("start")
                elapsed_value = status_element.get("elapsed")
                start_datetime = datetime.fromisoformat(starttime_value)
                # Convert the elapsed time to a float
                elapsed_seconds = float(elapsed_value)
                # Create a timedelta object with the elapsed time
                elapsed_timedelta = timedelta(seconds=elapsed_seconds)
                # Add the elapsed time to the start datetime
                end_datetime = start_datetime + elapsed_timedelta
                # Convert the result to a string
                endtime_value = end_datetime.isoformat()
                end_datetime = datetime.fromisoformat(endtime_value)
                start_time = start_datetime.replace(microsecond=0)
                end_time = end_datetime.replace(microsecond=0)
                time_difference = end_time - start_time
                minutes, seconds = divmod(time_difference.seconds, 60)
                formatted_time = f"{minutes} min {seconds} sec"
                # remove "-" from start and end time
                start_time = str(start_time).replace("-", "")
                end_time = str(end_time).replace("-", "")

                kw_elements = test_data.findall(".//kw[@name='Log Actual And Expected Results']")
                actual_results = []
                expected_results = []
                for kw_element in kw_elements:
                    kw_status_element = kw_element.find("status")
                    kw_status = kw_status_element.get("status")
                    if status == 'PASS':
                        msg_elements = kw_element.findall("kw/msg[@level='INFO']")
                        actual_result = msg_elements[0].text.strip() if len(msg_elements) >= 1 else None
                        expected_result = msg_elements[1].text.strip() if len(msg_elements) >= 2 else None
                        
                        actual_results.append(actual_result)
                        expected_results.append(expected_result)
                    elif status == 'NOT RUN' or status == 'FAIL' or status == 'SKIP':
                        msg_elements = kw_element.findall("./arg")
                        actual_result = "The expected result has not been met"
                        expected_result = msg_elements[1].text.strip() if len(msg_elements) >= 2 else None
                        # actual_result = 'Failed to '+ testcase_name
                        # expected_result = testcase_name
                        actual_results.append(actual_result)
                        expected_results.append(expected_result)

                actual_text = ' '.join(actual_results).strip() if any(actual_results) else "Expected Result is not Displayed"
                expected_text = ' '.join(expected_results).strip() if any(expected_results) else expected_results

                # Append data only if kw_status is not 'SKIP'
                if status != 'SKIP':
                    data_rows.append((suite_name, testcase_name, status, start_time, end_time, formatted_time, expected_text, actual_text))
                # suite_times.append(suite_time)
        return data_rows, suite_times

    def seconds_to_minutes_and_seconds(self, seconds):
        seconds = float(seconds)  # Convert to float (or int) if it's a string
        minutes = int(seconds // 60)
        remaining_seconds = int(seconds % 60)
        return f"{minutes} min {remaining_seconds} sec"

    def read_html_xml_and_write_html(self, url, xml_path, output_path, execution_mode='Remote', timestamp='False'):
        current_username = getpass.getuser()
        # Read the Output.XML file
        if execution_mode=='Remote':
            data_rows, suite_times = self.read_log_from_xml(xml_path)
            start_time, end_time, execution_time, passed, failed, skipped, build_no, total_tc, suite_statistic = self.read_xml(xml_path)
        else:
            data_rows, suite_times = self.read_log_from_local_xmlfile(xml_path)
            start_time, end_time, execution_time, passed, failed, skipped, build_no, total_tc, suite_statistic = self.read_local_xmlfile(xml_path)
        actual_total = int(passed)+int(failed)
        percentage = (int(passed)/int(actual_total))*100
        rounded_percentage = round(percentage, 2)
        if timestamp == 'True':
            output_html_path = output_path + build_no+".html"
        else:
            output_html_path = output_path +".html"

        # Pie chart, where the slices will be ordered and plotted counter-clockwise:
        labels = 'Fail', 'Pass'
        sizes = [failed, passed]
        colors = ['red', 'green']
        explode = (0, 0.02)
        # Plotting the pie chart
        plt.figure(figsize=(3, 3), facecolor='none')
        plt.pie(sizes, explode=explode, labels=labels, colors=colors, autopct='%1.1f%%', startangle=90, wedgeprops=dict(width=0.4, edgecolor='w'))
        plt.axis('equal')  # Equal aspect ratio ensures that pie is drawn as a circle.
        # Save the plot as an image file
        plt.savefig('chart.png', transparent=True)
        with open('chart.png', 'rb') as image_file:
            base64_image = base64.b64encode(image_file.read()).decode('utf-8')
        with open('Logo.png', 'rb') as image_file1:
            base64_image1 = base64.b64encode(image_file1.read()).decode('utf-8')

        if timestamp == 'True':
            splitted_xml_name = xml_path.split("put")
            timestamp = splitted_xml_name[1].replace(".xml", '')
            log_name = "log"+timestamp+".html"
        else:
            log_name = "log.html"
        result_text = 'FAILURE' if int(failed) > 0 else 'SUCCESS'
        result_color = 'failure' if int(failed) > 0 else 'success'
        with open(output_html_path, 'w', encoding='utf-8') as html_file:
            # Write print statements to a html file
            html_content = f"<html>\n<head>\n<style>\n"
            # html_content += "body { background-color: #f0f0f0; }\n"
            html_content += "strong.keys { font-weight: bold; color: black; }\n"
            html_content += "strong.header { font-weight: bold; font-size: 30px; color: green; }\n"
            html_content += "strong.glow { font-weight: bold; color: black; animation: glow 1s ease-in-out infinite alternate; }\n"
            html_content += "@keyframes glow { to { text-shadow: 0 0 20px #00fffc, 0 0 60px #00fffc; } }\n"
            html_content += "strong { font-weight: bold; color: blue; }\n"
            html_content += "strong.result.success { color: green; }"
            html_content += "strong.result.failure { color: red; }"
            html_content += "table { border-collapse: collapse; border: 1px solid black; width: 100%; }\n"
            html_content += "th { border: 1px solid black; text-align: center; padding: 8px; }\n"
            html_content += "td { border: 1px solid black; text-align: center; padding: 8px; }\n"
            html_content += "</style>\n"
            html_content += "<script>\n"
            html_content += "function openInNewTab(url) {\n"
            html_content += "    window.open(url, '_blank');\n"
            html_content += "}\n"
            html_content += "</script>\n"
            html_content += "</head>\n<body>"
            html_content += "<div style='display: flex;'>"
            html_content += "    <div style='flex: 1;'>"
            html_content += "       <div style='text-align: center;'><strong class='header'>VLMS Test Execution Report</strong></div>"
            html_content += "    </div>"
            html_content += "    <div style='flex: 1; text-align: right;'>"
            html_content += f"       <h><img src='data:image/png;base64,{base64_image1}' alt='Vg_logo'></h>"
            html_content += "    </div>"
            html_content += "</div>"                
            html_content += "<br>\n\n\n"
            html_content += "<div style='display: flex;'>"
            html_content += "    <div style='flex: 0.5;'>"
            html_content += "        <p>\n<strong class='keys glow'>Plan Report:</strong></p>"
            html_content += "        <p><strong>Name:</strong> Regression Test Suites</p>\n"
            html_content += f"        <p><strong>Triggered By:</strong> {current_username}</p>\n"
            html_content += f"        <p><strong>Environment Url:</strong> {url}</p>\n"
            html_content += f"        <p><strong>Start Time:</strong> {start_time}</p>\n"
            html_content += f"        <p><strong>End Time:</strong> {end_time}</p>\n"
            html_content += f"        <p><strong>Execution Time:</strong> {execution_time}</p>\n"
            html_content += f"        <p><strong>Result:</strong> <strong class='result {result_color}'> {result_text}</strong></p>\n"
            html_content += f"        <p><strong>Build Number:</strong> {build_no}</p>\n "
            html_content += "    </div>"
            html_content += "    <div style='flex: 0.3;'>"
            html_content += "<br>\n"
            # html_content += "        <h><strong class='keys'>Test Case Results</strong></h><p><img src='chart.png' alt='Test Case Pie Chart'></p>"
            html_content += f"       <h><strong class='keys glow'>Test Case Results</strong></h><p><img src='data:image/png;base64,{base64_image}' alt='Test Case Pie Chart'></p>"
            html_content += "    </div>"
            html_content += "</div>"
            html_content += "<p><strong class='keys glow'>Overall Summary:</strong></p>\n\n"
            html_content += "<table style='width: 40%;'>"
            html_content += "<tr><th style='width: 15%;'>Statistic</th><th style='width: 10%;'>Value</th></tr>"
            html_content += f"<tr><td><strong>Testcases Passed:</strong></td><td>{passed}</td></tr>"
            html_content += f"<tr><td><strong>Testcases Failed:</strong></td><td>{failed}</td></tr>"
            # html_content += f"<tr><td><strong>Testcases Retried:</strong></td><td>{skipped}</td></tr>"
            html_content += f"<tr><td><strong>Testcases Total:</strong></td><td>{actual_total}</td></tr>"
            html_content += f"<tr><td><strong>Pass Percentage:</strong></td><td><strong class='result success'>{rounded_percentage}%</strong></td></tr>"
            html_content += "</table>\n"
            html_content += "<br>\n\n\n"
            html_content += "<br>\n\n\n"
            # Write the table for Test and Suite statistics to the HTML file
            html_content += "<p><strong class='keys glow'>Suite Statistics:</strong></p>\n"
            html_content += "<table style='width: 80%;'>"
            html_content += "\n<tr><th>Sl. No</th><th style='width: 35%;'>Suite Name</th><th>Status</th><th>Total</th><th>Pass</th><th>Fail</th><th style='width: 20%;'>Duration</th></tr>\n"
            time_index = 0
            for row in suite_statistic:
                html_content += "<tr>\n"
                for col_index, value in enumerate(row):
                    # Apply different styles based on the column index
                        if col_index == 2:  # Status column
                            status_color = 'green' if value == 'PASS' else 'red' if value == 'FAIL' else 'orange'
                            html_content += f"<td style='color: {status_color};'>{value}</td>"
                        elif col_index == 1:  # Test Name column
                            html_content += f"<td onclick=\"openInNewTab('{log_name}')\" style='color: blue;'>{value}</td>"
                        elif col_index == 4:
                            html_content += f"<td style='color: green;'>{value}</td>"
                        elif col_index == 5:
                            html_content += f"<td style='color: red;'>{value}</td>"
                        else:
                            html_content += f"<td>{value}</td>"
                if time_index < len(suite_times):
                    html_content += f"<td>{suite_times[time_index]}</td>"
                else:
                    html_content += "<td></td>"  # Handle the case where suite_times is shorter than suite_statistic
                time_index += 1  # Increment the index for suite_times
                html_content += "</tr>\n"
            html_content += "</table>\n"
            html_content += "<br>\n\n"
            html_content += "<br>\n\n"
                
            html_content += "<p><strong class='keys glow'>Test Results:</strong></p>\n"
            html_content += "<table>\n<tr><th style='width: 15%;'>Suite Name</th><th style='width: 20%;'>Test Name</th><th style='width: 7%;'>Status</th><th style='width: 9%;'>Start Time (IST)</th><th style='width: 9%;'>End Time (IST)</th><th style='width: 8%;'>Duration</th><th style='width: 17%;'>Expected Result</th><th style='width: 17%;'>Actual Result</th></tr>\n"
            for row in data_rows:
                html_content += "<tr>\n"
                for index, value in enumerate(row):
                    # Apply different styles based on the column index
                    if index == 2:  # Status column
                        status_color = 'green' if value == 'PASS' else 'red' if value == 'FAIL' else 'orange'
                        html_content += f"<td style='color: {status_color};'>{value}</td>"
                    elif index == 1:  # Test Name column
                        html_content += f"<td onclick=\"openInNewTab('{log_name}')\" style='color: blue;'>{value}</td>"
                    elif index == 7:
                        status_color = 'red' if value.startswith('The expected result has not been met') else 'green'
                        html_content += f"<td style='color: {status_color};'>{value}</td>"
                    elif index == 6:
                        html_content += f"<td style='color: blue;'>{value}</td>"
                    else:
                        html_content += f"<td>{value}</td>"
                html_content += "</tr>\n"
            html_content += "</table>\n"
            html_content += "<br>\n\n"
            html_content += "<br>\n\n"
            
            html_content += "</body>\n</html>"

            html_file.write(html_content)
        os.remove('chart.png')
    
    def read_log_from_local_xmlfile(self, input_xml_path):
        tree = ET.parse(input_xml_path)
        root = tree.getroot()
        suite_datas = root.findall(".//suite[@id='s1']")
        if suite_datas:
            first_suite = suite_datas[0]
            execution_name = first_suite.get("name")
            print("Execution Name:", execution_name)
            if execution_name == 'VLMS4.2 Regression':
                suite_datas = first_suite.findall(".//suite")
        data_rows = []
        suite_times = []
        for suite_data in suite_datas:
            suite_name = suite_data.get("name")
            suite_status = suite_data.find("./status")
            suite_starttime_in_seconds = suite_status.get("starttime")
            suite_endtime_in_seconds = suite_status.get("endtime")
            time_format = "%Y%m%d %H:%M:%S.%f"
            suite_starttime = time.mktime(time.strptime(suite_starttime_in_seconds, time_format))
            suite_endtime = time.mktime(time.strptime(suite_endtime_in_seconds, time_format))
            time_dif = suite_endtime - suite_starttime
            suite_time = self.seconds_to_minutes_and_seconds(time_dif)
            suite_times.append(suite_time)
            test_datas = suite_data.findall("test")
            for test_data in test_datas:
                testcase_name = test_data.get("name")
                status_element = test_data.find("status")
                status = status_element.get("status")
                start_time_with_millis = status_element.get("starttime")
                parsed_timestamp = datetime.strptime(start_time_with_millis, "%Y%m%d %H:%M:%S.%f")
                starttime = parsed_timestamp.strftime("%Y%m%d %H:%M:%S")
                end_time_with_millis = status_element.get("endtime")
                parsed_timestamp1 = datetime.strptime(end_time_with_millis, "%Y%m%d %H:%M:%S.%f")
                endtime = parsed_timestamp1.strftime("%Y%m%d %H:%M:%S")
                time_format = "%Y%m%d %H:%M:%S"
                # Parse timestamps and convert to Unix timestamps
                start_time = time.mktime(time.strptime(starttime, time_format))
                end_time = time.mktime(time.strptime(endtime, time_format))
                time_difference = end_time - start_time
                minutes, seconds = divmod(int(time_difference), 60)
                formatted_time = f"{minutes} min {seconds} sec"

                kw_elements = test_data.findall(".//kw[@name='Log Actual And Expected Results']")
                actual_results = []
                expected_results = []
                for kw_element in kw_elements:
                    kw_status_element = kw_element.find("status")
                    kw_status = kw_status_element.get("status")
                    if status == 'PASS':
                        msg_elements = kw_element.findall("kw/msg[@level='INFO']")
                        actual_result = msg_elements[0].text.strip() if len(msg_elements) >= 1 else None
                        expected_result = msg_elements[1].text.strip() if len(msg_elements) >= 2 else None
                        
                        actual_results.append(actual_result)
                        expected_results.append(expected_result)
                    elif status == 'NOT RUN' or status == 'FAIL' or status == 'SKIP':
                        msg_elements = kw_element.findall("./arg")
                        actual_result = "The expected result has not been met"
                        expected_result = msg_elements[1].text.strip() if len(msg_elements) >= 2 else None
                        # actual_result = 'Failed to '+ testcase_name
                        # expected_result = testcase_name
                        actual_results.append(actual_result)
                        expected_results.append(expected_result)

                actual_text = ' '.join(actual_results).strip() if any(actual_results) else "The expected result has not been met"
                expected_text = ' '.join(expected_results).strip() if any(expected_results) else expected_results

                # Append data only if kw_status is not 'SKIP'
                if status != 'SKIP':
                    data_rows.append((suite_name, testcase_name, status, starttime, endtime, formatted_time, expected_text, actual_text))
        return data_rows, suite_times

    def read_local_xmlfile(self, input_xml_path):
            # Parse the XML data from the file
            tree = ET.parse(input_xml_path)
            root = tree.getroot()
            status_elements = root.findall(".//test/status")
            # Initialize a counter for 'SKIP'
            skip_count = 0
            # Iterate through the status elements
            for status_element in status_elements:
                status_value = status_element.get('status')
                if status_value == 'SKIP':
                    skip_count += 1
            # Print the count of 'SKIP'
            print("Number of 'SKIP' status:", skip_count)
            statistics = root.find(".//statistics/total/stat")
            passed = statistics.get('pass')
            failed = statistics.get('fail')
            total_tc = int(passed)+int(failed)+int(skip_count)
            suite_statistics = root.findall(".//statistics/suite/stat")
            suite_statistic = []
            for index, suite in enumerate(suite_statistics):
                if index == 0:
                    # Skip the first iteration
                    continue
                suite_name = suite.get('name')
                suite_pass = suite.get('pass')
                suite_fail = suite.get('fail')
                suite_total = int(suite_pass)+int(suite_fail)
                suite_status = 'FAIL' if int(suite_fail) > 0 else 'PASS'
                suite_statistic.append([index, suite_name, suite_status, suite_total, suite_pass, suite_fail])
            print(suite_statistic)       
            suite_with_stat = root.find(".//suite/status")
            if suite_with_stat is not None:
                start_time_with_millis = suite_with_stat.get('starttime')
                parsed_timestamp = datetime.strptime(start_time_with_millis, "%Y%m%d %H:%M:%S.%f")
                start_time = parsed_timestamp.strftime("%Y%m%d %H:%M:%S")
                end_time_with_millis = suite_with_stat.get('endtime')
                parsed_timestamp1 = datetime.strptime(end_time_with_millis, "%Y%m%d %H:%M:%S.%f")
                end_time = parsed_timestamp1.strftime("%Y%m%d %H:%M:%S")
                splitted_starttime = str(start_time).split(" ")
                date = splitted_starttime[0].replace("-", "")
                log_hr = splitted_starttime[1].split(":")[0]
                build_no = date+"."+log_hr
                print(build_no)
                # Adjust the format string to exclude milliseconds
                time_format = "%Y%m%d %H:%M:%S"
                # Parse timestamps and convert to Unix timestamps
                start_time_in_seconds = time.mktime(time.strptime(start_time, time_format))
                end_time_in_seconds = time.mktime(time.strptime(end_time, time_format))
                time_difference = end_time_in_seconds - start_time_in_seconds
                print(time_difference)
                minutes, seconds = divmod(int(time_difference), 60)
                execution_time = f"{minutes} min {seconds} sec"
                print(execution_time)
           
            return start_time, end_time, execution_time, passed, failed, skip_count, build_no, total_tc, suite_statistic
